import os
import pandas as pd
import sqlite3
import tempfile
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
import streamlit as st

# ==============================
# Save to SQLite
# ==============================
def save_to_sqlite(df, db_file="uploaded_data.db"):
    conn = sqlite3.connect(db_file)
    df.to_sql("uploaded_data", conn, if_exists="replace", index=False)
    conn.close()

def load_from_sqlite(db_file="uploaded_data.db"):
    conn = sqlite3.connect(db_file)
    df = pd.read_sql("SELECT * FROM uploaded_data", conn)
    conn.close()
    return df

# ==============================
# Generate PDFs without Word/LibreOffice
# ==============================
def generate_pdfs_no_libreoffice(template_file, df):
    pdf_paths = []
    for idx, row in df.iterrows():
        context = row.to_dict()

        # Load DOCX template
        doc = Document(template_file)

        # Replace placeholders {{col}}
        for p in doc.paragraphs:
            for key, val in context.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in p.text:
                    p.text = p.text.replace(placeholder, str(val))

        # Create PDF
        out_pdf = os.path.join(
            tempfile.gettempdir(),
            f"{context.get('OrderID', idx)}_{context.get('Name','Record')}.pdf"
        )
        c = canvas.Canvas(out_pdf, pagesize=A4)
        text = c.beginText(50, 800)

        # Write each line from DOCX
        for p in doc.paragraphs:
            text.textLine(p.text)

        c.drawText(text)
        c.save()
        pdf_paths.append(out_pdf)

    return pdf_paths

# ==============================
# Streamlit UI
# ==============================
st.set_page_config(page_title="Custom PDF Generator", page_icon="📄", layout="centered")

st.title("📄 Custom PDF Generator")
st.markdown(
    "Upload a **DOCX template** (with placeholders like `{{ColumnName}}`) and a **CSV/Excel file**. "
    "The app will fill in the template with each row of data and generate PDFs."
)

# Upload template and data
template_file = st.file_uploader("Upload DOCX Template", type=["docx"])
data_file = st.file_uploader("Upload CSV/Excel Data", type=["csv", "xlsx"])

if template_file and data_file:
    # Save template to temp file
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_tpl:
        tmp_tpl.write(template_file.read())
        template_path = tmp_tpl.name

    # Load data
    if data_file.name.endswith(".csv"):
        df = pd.read_csv(data_file)
    else:
        df = pd.read_excel(data_file)

    # Save to SQLite
    save_to_sqlite(df)

    st.write("### 📊 Preview of Uploaded Data")
    st.dataframe(df.head())

    if st.button("Generate PDFs"):
        pdf_paths = generate_pdfs_no_libreoffice(template_path, df)

        st.success("✅ PDFs generated successfully!")
        for pdf in pdf_paths:
            with open(pdf, "rb") as f:
                st.download_button(
                    label=f"⬇ Download {os.path.basename(pdf)}",
                    data=f,
                    file_name=os.path.basename(pdf),
                    mime="application/pdf"
                )
